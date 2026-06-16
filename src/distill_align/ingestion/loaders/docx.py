"""
DOCX file loader.

Handles loading and metadata extraction from Microsoft Word documents.
"""

from ...core.exceptions import LoaderError
from ...core.schemas import SourceMetadata
from .base import BaseLoader


class DOCXLoader(BaseLoader):
    """Loader for Microsoft Word (.docx) files."""

    SUPPORTED_EXTENSIONS = {".docx"}

    def load(self) -> str:
        """
        Load DOCX file content.

        Returns:
            Extracted text content as string.

        Raises:
            LoaderError: If file cannot be read.
        """
        try:
            from docx import Document  # type: ignore[import-not-found]

            doc = Document(str(self.file_path))
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            return "\n\n".join(paragraphs)

        except ImportError:
            raise LoaderError(
                "python-docx is required for DOCX loading. Install with: pip install python-docx"
            ) from None
        except Exception as e:
            raise LoaderError(f"Failed to read DOCX file: {e}") from e

    def extract_metadata(self) -> SourceMetadata:
        """
        Extract metadata from DOCX file.

        Returns:
            SourceMetadata with file information.
        """
        try:
            from docx import Document  # type: ignore[import-not-found]

            doc = Document(str(self.file_path))
            props = doc.core_properties

            return SourceMetadata(
                source_type="text",
                file_path=str(self.file_path),
                file_name=self.file_path.name,
                title=props.title or self.file_path.stem,
                author=props.author,
                custom_tags={
                    "format": "docx",
                    "subject": props.subject,
                    "keywords": str(props.keywords) if props.keywords else None,
                    "created": str(props.created) if props.created else None,
                    "modified": str(props.modified) if props.modified else None,
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                },
            )
        except ImportError:
            return SourceMetadata(
                source_type="text",
                file_path=str(self.file_path),
                file_name=self.file_path.name,
                title=self.file_path.stem,
                custom_tags={"format": "docx"},
            )
        except Exception as e:
            raise LoaderError(f"Failed to extract DOCX metadata: {e}") from e
